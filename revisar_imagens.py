import asyncio
import sys
import os
import time
import tempfile
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import subprocess
from playwright.sync_api import sync_playwright
from docx import Document

if sys.platform.startswith("win"):
    asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

# --- Configuração do Chrome Debug ---
CHROME_PATH = r"C:\Program Files\Google\Chrome\Application\chrome.exe"
CHROME_USER_DATA_DIR = r"C:\temp\chrome"
CHROME_REMOTE_DEBUGGING_PORT = 9222
CHROME_DEBUG_URL = f"http://localhost:{CHROME_REMOTE_DEBUGGING_PORT}"

# --- Funções reutilizadas ---
def salvar_texto_docx(respostas_dict, destino):
    doc = Document()
    for nome_arquivo, texto in respostas_dict.items():
        doc.add_heading(f"Arquivo: {nome_arquivo}", level=1)
        doc.add_paragraph(texto)
        doc.add_paragraph("\n")
    doc.save(destino)

def esperar_resposta_gpt(page, tempo_maximo=180, intervalo_check=1.5, tempo_estavel=4):
    conteudo_anterior = ""
    tentativas_estaveis = 0
    tempo_decorrido = 0
    while tempo_decorrido < tempo_maximo:
        if page.url.endswith("/api/auth/error"):
            return "__ERRO_AUTENTICACAO__"
        erro_elements = page.locator(".text-token-text-error")
        if erro_elements.count() > 0:
            return "__ERRO_GPT__"
        if page.locator("button:has(svg[aria-label='Stop generating'])").is_visible():
            time.sleep(intervalo_check)
            tempo_decorrido += intervalo_check
            continue
        elementos = page.locator(".markdown")
        respostas = elementos.all_text_contents()
        conteudo_atual = respostas[-1] if respostas else ""
        if conteudo_atual == conteudo_anterior:
            tentativas_estaveis += 1
        else:
            tentativas_estaveis = 0
        if tentativas_estaveis >= tempo_estavel:
            return conteudo_atual if conteudo_atual else "__ERRO_GPT__"
        conteudo_anterior = conteudo_atual
        time.sleep(intervalo_check)
        tempo_decorrido += intervalo_check
    return conteudo_anterior if conteudo_anterior else "__ERRO_GPT__"

def enviar_pdf_para_gpt(page, caminho_pdf):
    tempo_maximo = 900  # 15 minutos
    intervalo = 1.5
    tempo_decorrido = 0

    while page.locator("div[role='listitem']").is_visible() and tempo_decorrido < tempo_maximo:
        time.sleep(intervalo)
        tempo_decorrido += intervalo

    if page.locator("div[role='listitem']").is_visible():
        return "__ARQUIVO_JA_ANEXADO__"

    try:
        if page.locator("input[type='file']").count() > 0:
            page.set_input_files("input[type='file']", caminho_pdf)
        else:
            page.click("button:has(svg[aria-label='Upload a file'])", timeout=5000)
            page.wait_for_selector("input[type='file']", timeout=5000)
            page.set_input_files("input[type='file']", caminho_pdf)
    except:
        return "__ERRO_ENVIO__"

    time.sleep(4)
    page.keyboard.type("T1 - R2")
    page.keyboard.press("Enter")
    return esperar_resposta_gpt(page)

def abrir_chrome_debug():
    try:
        subprocess.Popen([
            CHROME_PATH,
            f"--remote-debugging-port={CHROME_REMOTE_DEBUGGING_PORT}",
            f"--user-data-dir={CHROME_USER_DATA_DIR}"
        ])
        messagebox.showinfo("Chrome", "Chrome com depuração iniciado.")
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao iniciar o Chrome:\n{e}")

def processar_pdfs():
    arquivos = filedialog.askopenfilenames(title="Selecione os arquivos PDF", filetypes=[("PDF files", "*.pdf")])
    if not arquivos:
        return

    respostas = {}
    texto_log.delete("1.0", tk.END)
    texto_log.insert(tk.END, "Iniciando processamento...\n")
    janela.update()

    with sync_playwright() as p:
        try:
            browser = p.chromium.connect_over_cdp(CHROME_DEBUG_URL)
            context = browser.contexts[0]
            page = context.pages[0] if context.pages else context.new_page()
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao conectar ao navegador Chrome com debug remoto:\n{e}")
            return

        for arquivo in arquivos:
            nome = os.path.basename(arquivo)
            texto_log.insert(tk.END, f"Processando: {nome}\n")
            janela.update()

            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                with open(arquivo, "rb") as f:
                    tmp.write(f.read())
                tmp_path = tmp.name

            resposta = enviar_pdf_para_gpt(page, tmp_path)

            if resposta.startswith("__ERRO"):
                texto_log.insert(tk.END, f"Erro ao processar {nome}: {resposta}\n")
            else:
                texto_log.insert(tk.END, f"{nome} processado com sucesso.\n")

            respostas[nome] = resposta
            janela.update()

    doc_path = os.path.join(os.getcwd(), "resultado.docx")
    salvar_texto_docx(respostas, doc_path)
    texto_log.insert(tk.END, f"\nProcessamento finalizado. Resultado salvo em: {doc_path}\n")
    janela.update()

# --- Interface Tkinter ---
janela = tk.Tk()
janela.title("Automatizador de PDFs para ChatGPT")
janela.geometry("600x400")

frame_botoes = tk.Frame(janela)
frame_botoes.pack(pady=10)

botao_abrir_chrome = tk.Button(frame_botoes, text="Abrir Chrome com Debug", command=abrir_chrome_debug)
botao_abrir_chrome.pack(side=tk.LEFT, padx=10)

botao_processar = tk.Button(frame_botoes, text="Selecionar e Processar PDFs", command=processar_pdfs)
botao_processar.pack(side=tk.LEFT, padx=10)

texto_log = scrolledtext.ScrolledText(janela, width=80, height=20)
texto_log.pack(pady=10)

janela.mainloop()
